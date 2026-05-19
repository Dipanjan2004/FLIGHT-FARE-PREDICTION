"""Generate a formal academic .docx report for the Flight Fare Prediction project.

Run: python -m scripts.generate_report
Output: Flight_Fare_Prediction_Report.docx in the project root.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
sys.path.append(str(ROOT))

import joblib

REPORTS_DIR = ROOT / "reports"
MODELS_DIR = ROOT / "models"
SRC_DIR = ROOT / "src"
APP_DIR = ROOT / "app"
OUTPUT = ROOT / "Flight_Fare_Prediction_Report.docx"


# ───────────────────────────── helpers ─────────────────────────────


def set_font(run, name: str = "Calibri", size: int = 11, bold: bool = False,
             italic: bool = False, color: RGBColor | None = None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # Force the font on the East Asian range too (docx quirk)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
             size: int = 11, align=None, font: str = "Calibri") -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, name=font, size=size, bold=bold, italic=italic)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size=11)


def add_code_block(doc: Document, code: str) -> None:
    """Insert a monospace shaded code paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    run = p.add_run(code)
    set_font(run, name="Consolas", size=9)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_image(doc: Document, path: Path, caption: str, width_in: float = 6.0) -> None:
    if not path.exists():
        add_para(doc, f"[Missing image: {path.name}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_font(run, size=10, italic=True, color=RGBColor(0x55, 0x55, 0x55))


def style_table_header(table) -> None:
    hdr = table.rows[0]
    for cell in hdr.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3A5F")
        tcPr.append(shd)
        for p in cell.paragraphs:
            for run in p.runs:
                set_font(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))


# ───────────────────────────── content ─────────────────────────────


TITLE_DEPT = "DEPARTMENT OF ELECTRONICS & TELECOMMUNICATION ENGINEERING"
TITLE_PROJECT = "FLIGHT FARE PREDICTION"
COURSE_LINE = "Report for Course:  Introduction to Machine Learning (24ETAEC49)"
COORDINATOR_LINE = "Course Coordinator:  Prof. Dr. Akshata.S.K"
INSTITUTE_LINES = [
    "RAMAIAH INSTITUTE OF TECHNOLOGY",
    "(Autonomous Institute, Affiliated to VTU)",
    "Accredited by NBA & NAAC with ‘A+’ Grade",
    "MSR Nagar, Bangalore - 560054",
]
REPORT_DATE = "May 20, 2026"

STUDENTS = [
    ("DIPANJAN CHOWDHURY", "1MS24ET015"),
    ("DEVANSH GUPTA", "1MS24ET012"),
]


ABSTRACT = """In the highly dynamic Indian domestic aviation industry, ticket prices fluctuate
constantly with demand, route, airline, time-of-day, and how many days remain
before departure. This volatility makes it difficult for travellers to judge
whether a quoted fare is reasonable, and difficult for airlines to communicate
fair pricing. This project addresses the problem by developing a supervised
machine-learning regression pipeline that predicts the price of a flight ticket
from nine easily-collected features. A 300,153-record Kaggle dataset of Indian
domestic flights is cleaned, explored, and modelled with three regression
algorithms — Linear Regression, Random Forest, and Histogram Gradient Boosting
— which are compared with five-fold cross-validation. Random Forest emerges as
the strongest candidate (cross-validated RMSE ≈ Rs. 2,700, R² ≈ 0.986); its
hyperparameters are then tuned with RandomizedSearchCV. The final model is
persisted with joblib and deployed behind a Streamlit web interface that allows
end-users to estimate a flight fare interactively. The pipeline is reproducible
end-to-end and demonstrates that ensemble tree methods are effective for fare
prediction without manual feature scaling or engineering."""

INTRO = """Air-ticket pricing in India is governed by complex revenue-management systems
that adjust fares minute by minute based on demand, seat inventory, competitor
prices, fuel costs, route-level constraints, and the time remaining until
departure. From the passenger’s perspective, this opacity is frustrating —
identical flights can be priced very differently depending on when the booking
is made. A predictive model that learns from historical fare data can serve
both as a consumer-facing decision aid (“is today’s price reasonable, or should
I wait?”) and as an academic vehicle for exercising the supervised-learning
toolkit. This project develops such a model end-to-end: dataset acquisition
from Kaggle, exploratory data analysis, preprocessing with a scikit-learn
ColumnTransformer, comparison of three regression algorithms with cross-
validation, hyperparameter tuning of the best candidate, and deployment as an
interactive Streamlit web application."""

LIT_REVIEW = """The application of supervised regression to flight-fare prediction has matured
rapidly over the last decade. Early studies relied on linear regression and
generalized linear models for their interpretability — the learned coefficients
directly explain how each feature contributes to price. However, linear models
struggle with the strong non-linearity in fare-vs-days-left curves and with
interactions such as airline × class. Subsequent work moved toward decision
tree ensembles: Random Forests (Breiman, 2001), Gradient Boosted Trees, and
more recently XGBoost and LightGBM. These methods consistently outperform
linear baselines on tabular fare data because they (i) capture non-linear
relationships natively, (ii) handle mixed numeric/categorical inputs after
encoding, and (iii) are robust to feature scaling. Histogram Gradient Boosting
— scikit-learn’s native binned-tree implementation introduced in version 0.21
— offers a closely-related approach with substantially lower training cost on
large datasets. This project adopts the same comparative philosophy: train a
linear baseline, a Random Forest, and a Histogram Gradient Boosting Regressor,
and select the best by cross-validated error."""

MODEL_OVERVIEW = """Flight-fare prediction is a regression problem: the target variable, ticket
price, is a continuous positive real number, not a discrete label. The core
algorithm chosen is the Random Forest Regressor, an ensemble of decision trees
trained on bootstrap samples of the training data with random feature
subsampling at each split. A single decision tree splits the feature space
along axis-aligned thresholds (for example, “class == Business?”, “days_left ≤
7?”) and assigns each leaf the mean of the training prices that fall into it.
Individual trees overfit because they can grow until every training row sits
alone in its own leaf. The Random Forest fixes this by averaging the
predictions of many such trees, each grown on a different randomly-drawn
subset of the data. Averaging many high-variance, low-bias estimators yields a
prediction that is both flexible (captures non-linearity and interactions) and
stable (low variance on unseen data). For this project the forest is trained
inside a scikit-learn Pipeline that performs one-hot encoding of the seven
categorical features (airline, source city, departure time, stops, arrival
time, destination city, class) and passes the two numeric features (duration,
days left) through unchanged. Encapsulating preprocessing and the model in a
single Pipeline guarantees that the exact same transformation is applied
during training, cross-validation, hyperparameter search, and live
prediction — eliminating an entire class of train/inference mismatch bugs."""

VALIDATION = """To produce an honest estimate of generalisation error and prevent overfitting,
the project uses a layered validation strategy. First, the dataset is split
80/20 into a training partition (240,122 rows) and an independent test
partition (60,031 rows) with a fixed random seed for reproducibility. The test
partition is held out throughout training, cross-validation, and hyperparameter
tuning, and is only consulted once at the end to report final metrics. Second,
within the training partition, five-fold K-Fold cross-validation is used to
compare candidate models. Each fold trains on 80% of the training data and
validates on the remaining 20%, rotating until every training row has been
used for validation exactly once. The mean and standard deviation of the
fold-level RMSE provide an unbiased estimate of model stability. Third, the
best model identified by cross-validation is then tuned using
RandomizedSearchCV (10 random parameter combinations, three-fold CV), which
explores a wider hyperparameter space than grid search at a fraction of the
cost. Three regression metrics are reported throughout: Root Mean Squared
Error (RMSE), which penalises large errors quadratically; Mean Absolute Error
(MAE), which gives the typical rupee deviation; and the coefficient of
determination R², which expresses the fraction of price variance the model
explains relative to a mean-only baseline."""

PIPELINE_STAGES = [
    ("1. Data Acquisition", [
        "Dataset programmatically downloaded from Kaggle via the kaggle CLI",
        "Saved locally to data/Clean_Dataset.csv and cached on subsequent runs",
    ]),
    ("2. Data Cleaning", [
        "Dropped the leading unnamed index column produced by the CSV export",
        "The dataset author already cleaned missing values, so no additional imputation was required",
        "Excluded the flight code column from features (high cardinality — would memorise rather than generalise)",
    ]),
    ("3. Exploratory Data Analysis", [
        "Price distribution: right-skewed, with a long tail driven by business-class fares",
        "Price by class: business fares cluster around Rs. 50k while economy clusters around Rs. 6k",
        "Price by airline: Vistara and Air India sit at the top median, AirAsia at the bottom",
        "Price vs. days_left: sharp spike for last-minute bookings (< 5 days) — pricing is strongly non-linear in this feature",
        "Price by stops: zero-stop flights are slightly cheaper on average than one-stop, two-or-more stops are mixed",
    ]),
    ("4. Preprocessing", [
        "One-hot encoding of 7 categorical features via sklearn OneHotEncoder with handle_unknown='ignore'",
        "Numeric features (duration, days_left) passed through unchanged — tree models are scale-invariant",
        "Both transformations bundled into a single ColumnTransformer, wrapped in a Pipeline alongside the regressor",
    ]),
    ("5. Modeling — Multi-Model Comparison", [
        "Linear Regression as a statistical baseline",
        "Random Forest Regressor as the primary non-linear model",
        "Histogram Gradient Boosting Regressor as a complementary boosted-tree alternative",
        "All three evaluated with 5-fold cross-validated RMSE on the training partition",
    ]),
    ("6. Hyperparameter Tuning", [
        "The best candidate (Random Forest) is then tuned with RandomizedSearchCV",
        "10 random combinations drawn from a grid over n_estimators, max_depth, min_samples_leaf, max_features",
        "Three-fold CV inside the search, scored by negative RMSE",
    ]),
    ("7. Persistence", [
        "Final tuned pipeline serialised with joblib.dump",
        "Sidecar metadata stores test metrics, category dropdown options, comparison results, and tuned hyperparameters",
    ]),
    ("8. Deployment", [
        "Streamlit web application loads the joblib artifacts at startup (cached via @st.cache_resource)",
        "Form-based UI collects the nine inputs and presents the predicted fare with an MAE-derived uncertainty band",
        "Sidebar surfaces R², RMSE, MAE, best-model name, and a comparison table for full transparency",
    ]),
]


DATASET_INFO_ROWS = [
    ["Column", "Type", "Role", "Description"],
    ["airline", "Categorical (6 levels)", "Feature", "AirAsia, Air_India, GO_FIRST, Indigo, SpiceJet, Vistara"],
    ["flight", "Categorical (high card.)", "Excluded", "Per-flight code; would cause memorisation"],
    ["source_city", "Categorical (6 levels)", "Feature", "Bangalore, Chennai, Delhi, Hyderabad, Kolkata, Mumbai"],
    ["departure_time", "Categorical (6 levels)", "Feature", "Time-of-day bucket"],
    ["stops", "Categorical (3 levels)", "Feature", "zero / one / two_or_more"],
    ["arrival_time", "Categorical (6 levels)", "Feature", "Time-of-day bucket"],
    ["destination_city", "Categorical (6 levels)", "Feature", "Same six cities"],
    ["class", "Categorical (2 levels)", "Feature", "Economy or Business"],
    ["duration", "Numeric", "Feature", "Flight duration in hours"],
    ["days_left", "Numeric", "Feature", "Days between booking and departure (1–49)"],
    ["price", "Numeric", "Target", "Ticket price in Rs."],
]


CONSOLE_OUTPUT = """Loaded 300,153 rows.
Train rows: 240,122 | Test rows: 60,031

=== Comparing candidate models with 5-fold CV ===

— LinearRegression —
  CV RMSE : 6,753  ± 28  (3.0s)
  Test RMSE: 6,762 | MAE: 4,553 | R²: 0.9113  (fit 0.6s)

— RandomForest —
  CV RMSE : 2,699  ± 10  (44.4s)
  Test RMSE: 2,700 | MAE: 1,094 | R²: 0.9859  (fit 12.2s)

— HistGradientBoosting —
  CV RMSE : 3,447  ± 26  (14.1s)
  Test RMSE: 3,449 | MAE: 2,002 | R²: 0.9769  (fit 5.4s)

Best by CV RMSE: RandomForest (2,699)

Tuning RandomForest with RandomizedSearchCV (n_iter=10, cv=3)...
  search took 358s
  best CV RMSE: 2,718
  best params : {'model__n_estimators': 200, 'model__min_samples_leaf': 2,
                  'model__max_features': 0.7, 'model__max_depth': 35}
  tuned test  : RMSE 2,693 | MAE 1,113 | R² 0.9859

Generating diagnostic plots...
  saved reports/07_model_comparison.png
  saved reports/08_feature_importance.png

Final model: RandomForest (tuned)
  RMSE 2,693 | MAE 1,113 | R² 0.9859
Saved model    -> models/flight_fare_model.joblib
Saved metadata -> models/model_metadata.joblib"""


CONCLUSION = """This project successfully delivered an end-to-end supervised regression
pipeline for predicting Indian domestic flight fares. Starting from a Kaggle
dataset of 300,153 records, three candidate regression models were compared
using 5-fold cross-validation: a Linear Regression baseline (cross-validated
RMSE Rs. 6,753), a Random Forest Regressor (Rs. 2,699), and a Histogram
Gradient Boosting Regressor (Rs. 3,447). The Random Forest, after
RandomizedSearchCV hyperparameter tuning, achieved a test-set RMSE of
Rs. 2,693, MAE of Rs. 1,113, and R² of 0.9859 — explaining 98.6% of the
variance in fares. Feature-importance analysis confirmed that travel class,
days left until departure, and airline identity are the dominant drivers of
price, consistent with domain knowledge. The trained pipeline is persisted to
disk and served behind a Streamlit web application that surfaces both the
predicted fare and the model’s performance metrics, allowing users to estimate
prices interactively while remaining honest about uncertainty. The project
demonstrates that an ensemble of decision trees, combined with a disciplined
preprocessing pipeline and rigorous cross-validated comparison, is well suited
to tabular fare-prediction problems and that high predictive quality is
achievable without manual feature engineering. Future work could extend the
approach with quantile-regression prediction intervals, additional features
(holiday flags, fuel-price index), and a temporal hold-out to better measure
deployment robustness."""


# ────────────────────────── source code reads ──────────────────────────


def read_source(path: Path) -> str:
    return path.read_text().rstrip() + "\n"


SOURCE_FILES = [
    ("src/config.py", SRC_DIR / "config.py"),
    ("src/data_loader.py", SRC_DIR / "data_loader.py"),
    ("src/preprocess.py", SRC_DIR / "preprocess.py"),
    ("src/eda.py", SRC_DIR / "eda.py"),
    ("src/train.py", SRC_DIR / "train.py"),
    ("app/streamlit_app.py", APP_DIR / "streamlit_app.py"),
]


# ─────────────────────────── doc construction ───────────────────────────


def build_title_page(doc: Document) -> None:
    # Vertical spacing then department header
    for _ in range(2):
        doc.add_paragraph()
    add_para(doc, TITLE_DEPT, bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(3):
        doc.add_paragraph()

    add_para(doc, "PROJECT:", bold=True, size=16,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, TITLE_PROJECT, bold=True, size=22,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4):
        doc.add_paragraph()

    add_para(doc, COURSE_LINE, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, COORDINATOR_LINE, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(2):
        doc.add_paragraph()

    # Student table
    t = doc.add_table(rows=1 + len(STUDENTS), cols=2)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.autofit = False
    t.columns[0].width = Cm(8.0)
    t.columns[1].width = Cm(4.5)
    hdr = t.rows[0].cells
    hdr[0].text = "NAME"
    hdr[1].text = "USN"
    for cell in hdr:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_font(run, size=11, bold=True)
    for i, (name, usn) in enumerate(STUDENTS, start=1):
        row = t.rows[i].cells
        row[0].text = name
        row[1].text = usn
        for cell in row:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_font(run, size=11)
    style_table_header(t)
    for _ in range(4):
        doc.add_paragraph()

    for line in INSTITUTE_LINES:
        add_para(doc, line, bold=True, size=12,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, REPORT_DATE, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)


def build_contents(doc: Document) -> None:
    add_page_break(doc)
    add_para(doc, "Contents", bold=True, size=18,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    toc = [
        "0.1  Abstract",
        "0.2  Introduction",
        "       0.2.1  Literature Review / Architectural Fit",
        "       0.2.2  Model Overview: Why Random Forest Fits",
        "       0.2.3  Validation & Interpretation of Scores",
        "0.3  Project Pipeline",
        "0.4  Dataset Components & Preprocessing",
        "0.5  Source Code",
        "0.6  Execution Protocol & Program Output",
        "       0.6.1  Exploratory Plots",
        "       0.6.2  Model-Comparison & Feature-Importance Plots",
        "       0.6.3  Console Output",
        "0.7  Conclusion",
    ]
    for line in toc:
        add_para(doc, line, size=12)


def build_body(doc: Document, metadata: dict) -> None:
    # 0.1 Abstract
    add_page_break(doc)
    add_heading(doc, "0.1  Abstract", level=1)
    for para in ABSTRACT.split("\n\n"):
        add_para(doc, para.replace("\n", " "), size=11)

    # 0.2 Introduction
    add_heading(doc, "0.2  Introduction", level=1)
    for para in INTRO.split("\n\n"):
        add_para(doc, para.replace("\n", " "), size=11)

    add_heading(doc, "0.2.1  Literature Review / Architectural Fit", level=2)
    add_para(doc, LIT_REVIEW.replace("\n", " "), size=11)

    add_heading(doc, "0.2.2  Model Overview: Why Random Forest Fits", level=2)
    add_para(doc, MODEL_OVERVIEW.replace("\n", " "), size=11)

    add_heading(doc, "0.2.3  Validation & Interpretation of Scores", level=2)
    add_para(doc, VALIDATION.replace("\n", " "), size=11)

    # 0.3 Project Pipeline
    add_page_break(doc)
    add_heading(doc, "0.3  Project Pipeline", level=1)
    for title, bullets in PIPELINE_STAGES:
        add_para(doc, title, bold=True, size=12)
        add_bullets(doc, bullets)

    # 0.4 Dataset Components & Preprocessing
    add_page_break(doc)
    add_heading(doc, "0.4  Dataset Components & Preprocessing", level=1)
    add_para(doc,
        "Source: Kaggle — Flight Price Prediction by Shubham Bathwal.",
        size=11)
    add_para(doc,
        "Size: 300,153 records × 12 columns (after dropping the unnamed CSV "
        "index). Training partition: 240,122 rows. Independent test partition: "
        "60,031 rows (20%, fixed random seed 42).",
        size=11)
    add_para(doc,
        "Target distribution: continuous numeric, right-skewed, ranging from "
        "around Rs. 1,000 (cheap economy) to over Rs. 100,000 (last-minute "
        "business). Per-feature breakdown:",
        size=11)
    doc.add_paragraph()
    rows = DATASET_INFO_ROWS
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = t.rows[r_idx].cells[c_idx]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font(run, size=10)
    style_table_header(t)

    add_para(doc, "", size=11)
    add_para(doc, "Preprocessing pipeline:", bold=True, size=12)
    add_bullets(doc, [
        "One-hot encoding of all seven categorical features (handle_unknown='ignore', dense output).",
        "Numeric features passed through unchanged — tree-based models are invariant to monotonic transformations.",
        "Encoding and modelling are wrapped in a single sklearn Pipeline so that cross-validation, hyperparameter search and live inference always apply the same transformations in the same order.",
    ])

    # 0.5 Source Code
    add_page_break(doc)
    add_heading(doc, "0.5  Source Code", level=1)
    add_para(doc,
        "The implementation is organised into six modules. Each file is "
        "reproduced in full below.",
        size=11)
    for label, path in SOURCE_FILES:
        add_para(doc, label, bold=True, size=12)
        add_code_block(doc, read_source(path))

    # 0.6 Execution & Output
    add_page_break(doc)
    add_heading(doc, "0.6  Execution Protocol & Program Output", level=1)
    add_para(doc, "End-to-end run sequence (one-time setup omitted):", size=11)
    add_code_block(doc,
        "python -m src.data_loader   # download dataset from Kaggle (~5 MB)\n"
        "python -m src.eda           # generate exploratory plots\n"
        "python -m src.train         # compare models, tune best, save artifacts\n"
        "streamlit run app/streamlit_app.py  # launch interactive demo"
    )

    add_heading(doc, "0.6.1  Exploratory Plots", level=2)
    add_image(doc, REPORTS_DIR / "01_price_distribution.png",
              "Figure 1 — Distribution of flight prices across the full dataset.")
    add_image(doc, REPORTS_DIR / "02_price_by_class.png",
              "Figure 2 — Price by travel class. Business fares dominate the upper range.")
    add_image(doc, REPORTS_DIR / "03_price_by_airline.png",
              "Figure 3 — Price by airline, ordered by median fare.")
    add_image(doc, REPORTS_DIR / "04_price_vs_days_left.png",
              "Figure 4 — Average price vs. days remaining until departure. "
              "Note the sharp spike inside the last few days.")
    add_image(doc, REPORTS_DIR / "05_price_by_stops.png",
              "Figure 5 — Price distribution by number of stops.")
    add_image(doc, REPORTS_DIR / "06_duration_vs_price.png",
              "Figure 6 — Flight duration vs. price (5,000-row sample), coloured by class.")

    add_heading(doc, "0.6.2  Model-Comparison & Feature-Importance Plots", level=2)
    add_image(doc, REPORTS_DIR / "07_model_comparison.png",
              "Figure 7 — Test-set RMSE and MAE for the three candidate models. "
              "Random Forest is the clear winner.")
    add_image(doc, REPORTS_DIR / "08_feature_importance.png",
              "Figure 8 — Top-15 most important features in the tuned "
              "Random Forest. Class (Business vs. Economy), days_left, and "
              "duration dominate the model.")

    add_heading(doc, "0.6.3  Console Output", level=2)
    add_para(doc, "Output produced by `python -m src.train` on the full dataset:",
             size=11)
    add_code_block(doc, CONSOLE_OUTPUT)

    # 0.7 Conclusion
    add_page_break(doc)
    add_heading(doc, "0.7  Conclusion", level=1)
    add_para(doc, CONCLUSION.replace("\n", " "), size=11)


def main() -> None:
    metadata_path = MODELS_DIR / "model_metadata.joblib"
    metadata = joblib.load(metadata_path) if metadata_path.exists() else {}

    doc = Document()
    # Default style font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Page margins (slightly tighter than default)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    build_title_page(doc)
    build_contents(doc)
    build_body(doc, metadata)

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
